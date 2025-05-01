"""
Servicio unificado para extracción y procesamiento de documentos.
"""

import logging
import os
import tempfile
import mimetypes
from datetime import datetime
from typing import Dict, Any, Optional, List, Callable, Protocol, Type
import asyncio

from fastapi import UploadFile

from common.config import get_settings
from common.db.storage import get_file_from_storage
from common.errors import DocumentProcessingError, ValidationError, ServiceError, ErrorCode, handle_errors
from common.context.vars import get_full_context
from common.context import with_context, Context
from common.cache import CacheManager, get_with_cache_aside, serialize_for_cache

import fitz  # PyMuPDF
from bs4 import BeautifulSoup
from docx import Document
import pandas as pd

logger = logging.getLogger(__name__)
settings = get_settings()

# --------------------------
# Clases de Estrategia para Extracción
# --------------------------

class DocumentExtractor(Protocol):
    """Protocolo que define la interfaz para extractores de documentos."""
    
    @staticmethod
    def can_handle(mimetype: str) -> bool:
        """Determina si este extractor puede manejar el tipo de documento dado."""
        ...
    
    @staticmethod
    async def extract(file_path: str) -> str:
        """Extrae texto del documento."""
        ...

class PDFExtractor:
    """Extractor para documentos PDF."""
    
    SUPPORTED_MIMETYPES = ['application/pdf']
    
    @staticmethod
    def can_handle(mimetype: str) -> bool:
        return mimetype in PDFExtractor.SUPPORTED_MIMETYPES
    
    @staticmethod
    async def extract(file_path: str) -> str:
        """Extrae texto de un archivo PDF."""
        try:
            # Evaluar tamaño del archivo antes de abrirlo
            file_size = os.path.getsize(file_path)
            page_count = PDFExtractor._get_page_count(file_path)
            
            # Determinar si es un PDF grande basado en tamaño o número de páginas
            large_file = page_count > 50 or file_size > 10 * 1024 * 1024
            
            if large_file:
                # Para archivos grandes, usar extracción optimizada
                return await LargePDFExtractor.extract(file_path)
            
            # Para archivos pequeños, usar extracción estándar pero controlada
            doc = None
            text = ""
            try:
                doc = fitz.open(file_path)
                for page in doc:
                    page_text = page.get_text("text")
                    if page_text.strip():
                        text += page_text + "\n\n"
                        
                # Si no se extrajo texto, intentar con otro modo
                if not text.strip() and doc:
                    for page in doc:
                        page_text = page.get_text("html")
                        soup = BeautifulSoup(page_text, "html.parser")
                        text_content = soup.get_text(separator=" ", strip=True)
                        if text_content.strip():
                            text += text_content + "\n\n"
            finally:
                # Asegurar que el documento siempre se cierre para liberar memoria
                if doc:
                    doc.close()
                    del doc
            
            return text.strip()
        except Exception as e:
            context = get_full_context()
            error_context = {
                "service": "ingestion",
                "operation": "extract_pdf",
                "file_path": file_path,
                **({k: v for k, v in context.items() if k in ["tenant_id", "collection_id"]} if context else {})
            }
            logger.error(f"Error extrayendo texto de PDF: {str(e)}", 
                        extra=error_context, exc_info=True)
            raise DocumentProcessingError(
                message=f"Error extracting text from PDF: {str(e)}",
                details={"file_path": file_path, "error_code": ErrorCode.DOCUMENT_EXTRACTION_ERROR},
                context=error_context
            )
    
    @staticmethod
    def _get_page_count(file_path: str) -> int:
        """Obtiene el número de páginas sin cargar todo el documento en memoria."""
        doc = None
        try:
            doc = fitz.open(file_path)
            return doc.page_count
        except Exception as e:
            context = get_full_context()
            error_context = {
                "service": "ingestion",
                "operation": "get_page_count",
                "file_path": file_path,
                **({k: v for k, v in context.items() if k in ["tenant_id", "collection_id"]} if context else {})
            }
            logger.warning(f"No se pudo obtener el número de páginas: {str(e)}", 
                          extra=error_context)
            return 0
        finally:
            # Garantizar que se liberen los recursos
            if doc:
                doc.close()
                del doc

class LargePDFExtractor:
    """Extractor especializado para PDFs grandes."""
    
    SUPPORTED_MIMETYPES = ['application/pdf']
    # Tamaño de lote óptimo para procesamiento de páginas
    BATCH_SIZE = 10
    
    @staticmethod
    def can_handle(mimetype: str) -> bool:
        return mimetype in LargePDFExtractor.SUPPORTED_MIMETYPES
    
    @staticmethod
    async def extract(file_path: str) -> str:
        """Extrae texto de un archivo PDF grande usando procesamiento por partes."""
        try:
            # Obtener información sin cargar todo el documento
            page_count = PDFExtractor._get_page_count(file_path)
            if page_count == 0:
                context = get_full_context()
                error_context = {
                    "service": "ingestion",
                    "operation": "extract_large_pdf",
                    "file_path": file_path,
                    **({k: v for k, v in context.items() if k in ["tenant_id", "collection_id"]} if context else {})
                }
                logger.error(f"No se pudo determinar el número de páginas para el PDF grande", 
                            extra=error_context)
                raise DocumentProcessingError(
                    message="Could not determine page count for large PDF",
                    details={"file_path": file_path},
                    context=error_context
                )
            
            # Crear lotes de páginas para procesamiento incremental
            batch_size = LargePDFExtractor.BATCH_SIZE
            batches = [(i, min(i + batch_size, page_count)) 
                      for i in range(0, page_count, batch_size)]
            
            all_text = []
            
            # Procesar cada lote independientemente para liberar memoria entre lotes
            for start_page, end_page in batches:
                batch_text = await LargePDFExtractor._process_page_range(
                    file_path, start_page, end_page
                )
                all_text.append(batch_text)
                
                # Forzar la liberación de memoria después de cada lote
                import gc
                gc.collect()
            
            return "\n\n".join(all_text).strip()
        except Exception as e:
            context = get_full_context()
            error_context = {
                "service": "ingestion",
                "operation": "extract_large_pdf",
                "file_path": file_path,
                **({k: v for k, v in context.items() if k in ["tenant_id", "collection_id"]} if context else {})
            }
            logger.error(f"Error extrayendo texto de PDF grande: {str(e)}", 
                        extra=error_context, exc_info=True)
            raise DocumentProcessingError(
                message=f"Error extracting text from large PDF: {str(e)}",
                details={"file_path": file_path, "error_code": ErrorCode.DOCUMENT_EXTRACTION_ERROR},
                context=error_context
            )
    
    @staticmethod
    async def _process_page_range(file_path: str, start_page: int, end_page: int) -> str:
        """Procesa un rango específico de páginas de un PDF."""
        doc = None
        try:
            doc = fitz.open(file_path)
            
            # Verificar que el documento está en el rango esperado
            if doc.page_count < end_page:
                end_page = doc.page_count
                
            batch_text = ""
            for i in range(start_page, end_page):
                try:
                    page = doc[i]
                    page_text = page.get_text("text")
                    if page_text.strip():
                        batch_text += page_text + "\n\n"
                    else:
                        # Intentar modo alternativo si no hay texto
                        page_text = page.get_text("html")
                        soup = BeautifulSoup(page_text, "html.parser")
                        text_content = soup.get_text(separator=" ", strip=True)
                        if text_content.strip():
                            batch_text += text_content + "\n\n"
                except Exception as e:
                    context = get_full_context()
                    error_context = {
                        "service": "ingestion",
                        "operation": "process_page_range",
                        "file_path": file_path,
                        "page": i,
                        **({k: v for k, v in context.items() if k in ["tenant_id", "collection_id"]} if context else {})
                    }
                    logger.warning(f"Error en página {i}: {str(e)}", 
                                  extra=error_context)
            
            return batch_text
        finally:
            # Garantizar que se liberen los recursos
            if doc:
                doc.close()
                del doc

class DocxExtractor:
    """Extractor para documentos Word (.docx)."""
    
    SUPPORTED_MIMETYPES = [
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword'
    ]
    
    @staticmethod
    def can_handle(mimetype: str) -> bool:
        return mimetype in DocxExtractor.SUPPORTED_MIMETYPES
    
    @staticmethod
    async def extract(file_path: str) -> str:
        """Extrae texto de un archivo Word (.docx)."""
        try:
            doc = Document(file_path)
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Procesar también tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            return "\n\n".join(paragraphs)
        except Exception as e:
            context = get_full_context()
            error_context = {
                "service": "ingestion",
                "operation": "extract_docx",
                "file_path": file_path,
                **({k: v for k, v in context.items() if k in ["tenant_id", "collection_id"]} if context else {})
            }
            logger.error(f"Error extrayendo texto de DOCX: {str(e)}", 
                        extra=error_context, exc_info=True)
            raise DocumentProcessingError(
                message=f"Error extracting text from DOCX: {str(e)}",
                details={"file_path": file_path},
                context=error_context
            )

class ExcelExtractor:
    """Extractor para hojas de cálculo (Excel, CSV)."""
    
    SUPPORTED_MIMETYPES = [
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'application/vnd.ms-excel',
        'text/csv',
        'application/csv'
    ]
    
    @staticmethod
    def can_handle(mimetype: str) -> bool:
        return mimetype in ExcelExtractor.SUPPORTED_MIMETYPES
    
    @staticmethod
    async def extract(file_path: str) -> str:
        """Extrae texto de un archivo Excel (.xlsx, .xls) o CSV."""
        try:
            if file_path.lower().endswith('.csv'):
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path, sheet_name=None)  # Leer todas las hojas
            
            text_parts = []
            
            # Si es un diccionario de DataFrames (múltiples hojas)
            if isinstance(df, dict):
                for sheet_name, sheet_df in df.items():
                    text_parts.append(f"Sheet: {sheet_name}")
                    text_parts.append(sheet_df.to_string(index=False))
            else:
                # Si es un único DataFrame
                text_parts.append(df.to_string(index=False))
            
            return "\n\n".join(text_parts)
        except Exception as e:
            context = get_full_context()
            error_context = {
                "service": "ingestion",
                "operation": "extract_excel",
                "file_path": file_path,
                **({k: v for k, v in context.items() if k in ["tenant_id", "collection_id"]} if context else {})
            }
            logger.error(f"Error extrayendo texto de Excel/CSV: {str(e)}", 
                        extra=error_context, exc_info=True)
            raise DocumentProcessingError(
                message=f"Error extracting text from Excel/CSV: {str(e)}",
                details={"file_path": file_path},
                context=error_context
            )

class HtmlExtractor:
    """Extractor para documentos HTML."""
    
    SUPPORTED_MIMETYPES = ['text/html', 'application/xhtml+xml']
    
    @staticmethod
    def can_handle(mimetype: str) -> bool:
        return mimetype in HtmlExtractor.SUPPORTED_MIMETYPES
    
    @staticmethod
    async def extract(file_path: str) -> str:
        """Extrae texto de un archivo HTML."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Eliminar scripts y estilos
            for script in soup(["script", "style"]):
                script.extract()
            
            # Obtener texto
            text = soup.get_text(separator=' ', strip=True)
            
            # Limpiar espacios excesivos
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
        except Exception as e:
            context = get_full_context()
            error_context = {
                "service": "ingestion",
                "operation": "extract_html",
                "file_path": file_path,
                **({k: v for k, v in context.items() if k in ["tenant_id", "collection_id"]} if context else {})
            }
            logger.error(f"Error extrayendo texto de HTML: {str(e)}", 
                        extra=error_context, exc_info=True)
            raise DocumentProcessingError(
                message=f"Error extracting text from HTML: {str(e)}",
                details={"file_path": file_path},
                context=error_context
            )

class TextExtractor:
    """Extractor para archivos de texto plano."""
    
    SUPPORTED_MIMETYPES = ['text/plain', 'text/markdown']
    
    @staticmethod
    def can_handle(mimetype: str) -> bool:
        return mimetype in TextExtractor.SUPPORTED_MIMETYPES
    
    @staticmethod
    async def extract(file_path: str) -> str:
        """Extrae texto de un archivo de texto plano."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            context = get_full_context()
            error_context = {
                "service": "ingestion",
                "operation": "extract_text",
                "file_path": file_path,
                **({k: v for k, v in context.items() if k in ["tenant_id", "collection_id"]} if context else {})
            }
            logger.error(f"Error leyendo archivo de texto: {str(e)}", 
                        extra=error_context, exc_info=True)
            raise DocumentProcessingError(
                message=f"Error reading text file: {str(e)}",
                details={"file_path": file_path},
                context=error_context
            )

# --------------------------
# Registro y Factory de Extractores
# --------------------------

# Lista de todos los extractores disponibles
DOCUMENT_EXTRACTORS = [
    PDFExtractor,
    DocxExtractor,
    ExcelExtractor,
    HtmlExtractor,
    TextExtractor
]

def get_extractor_for_mimetype(mimetype: str) -> Type[DocumentExtractor]:
    """
    Devuelve el extractor adecuado para un tipo MIME dado.
    
    Args:
        mimetype: Tipo MIME del documento
        
    Returns:
        Type[DocumentExtractor]: Clase del extractor
        
    Raises:
        DocumentProcessingError: Si no se encuentra un extractor para el tipo MIME
    """
    for extractor in DOCUMENT_EXTRACTORS:
        if extractor.can_handle(mimetype):
            return extractor
    
    context = get_full_context()
    error_context = {
        "service": "ingestion",
        "operation": "get_extractor",
        "mimetype": mimetype,
        **({k: v for k, v in context.items() if k in ["tenant_id", "collection_id"]} if context else {})
    }
    logger.error(f"No se encontró extractor para el tipo MIME: {mimetype}", 
                extra=error_context)
    raise DocumentProcessingError(
        message=f"No extractor found for MIME type: {str(mimetype)}",
        details={"mimetype": mimetype},
        context=error_context
    )

# --------------------------
# Funciones Principales de Extracción
# --------------------------

@with_context(tenant=False)
@handle_errors(error_type="service", log_traceback=True)
async def extract_text_from_file(file_path: str, mimetype: Optional[str] = None, ctx: Context = None) -> str:
    """
    Extrae texto de un archivo utilizando el extractor adecuado según su tipo MIME.
    
    Args:
        file_path: Ruta al archivo
        mimetype: Tipo MIME del archivo (opcional, se detecta si no se proporciona)
        
    Returns:
        str: Texto extraído del archivo
        
    Raises:
        DocumentProcessingError: Si hay un error en la extracción del texto
        ValidationError: Si el tipo de archivo no es soportado
    """
    # Obtener contexto para errores
    context = get_full_context()
    error_context = {
        "service": "ingestion",
        "operation": "extract_text",
        "file_path": file_path,
        **({k: v for k, v in context.items() if k in ["tenant_id", "collection_id"]} if context else {})
    }
    
    try:
        # Detectar tipo MIME si no se proporciona
        if not mimetype:
            mimetype, _ = mimetypes.guess_type(file_path)
            if not mimetype:
                raise ValidationError(
                    message="Could not determine file type",
                    details={"file_path": file_path, "error_code": ErrorCode.INVALID_FILE_TYPE},
                    context=error_context
                )
                
        # Obtener el extractor adecuado
        extractor_class = get_extractor_for_mimetype(mimetype)
        if not extractor_class:
            logger.error(f"No hay extractor disponible para el tipo {mimetype}", 
                        extra=error_context)
            raise ValidationError(
                message=f"Unsupported file type: {mimetype}",
                details={"mimetype": mimetype, "error_code": ErrorCode.UNSUPPORTED_FILE_TYPE},
                context=error_context
            )
        
        # Realizar la extracción
        logger.info(f"Iniciando extracción de texto para archivo tipo {mimetype}", 
                   extra=error_context)
        text = await extractor_class.extract(file_path)
        
        # Verificar resultado
        if not text or not text.strip():
            logger.warning(f"No se pudo extraer texto del archivo, intentando método alternativo",
                         extra=error_context)
            # Intentar extracción alternativa
            text = await extract_with_specific_method(file_path, mimetype)
            
        logger.info(f"Extracción completada: {len(text)} caracteres extraídos", 
                  extra=error_context)
        return text
    except (DocumentProcessingError, ValidationError):
        # Reenviar excepciones tipadas
        raise
    except Exception as e:
        # Capturar otras excepciones y convertirlas al formato estándar
        error_msg = f"Error inesperado extrayendo texto: {str(e)}"
        logger.error(error_msg, extra=error_context, exc_info=True)
        raise ServiceError(
            message=error_msg,
            details={"file_path": file_path, "error_code": ErrorCode.DOCUMENT_EXTRACTION_ERROR},
            context=error_context
        )

async def extract_with_specific_method(file_path: str, mimetype: str) -> str:
    """
    Intenta extraer texto con métodos específicos según el tipo MIME si el extractor principal falla.
    
    Args:
        file_path: Ruta al archivo
        mimetype: Tipo MIME del archivo
        
    Returns:
        str: Texto extraído del archivo
        
    Raises:
        DocumentProcessingError: Si hay un error en la extracción alternativa
    """
    context = get_full_context()
    error_context = {
        "service": "ingestion",
        "operation": "extract_with_specific_method",
        "file_path": file_path,
        **({k: v for k, v in context.items() if k in ["tenant_id", "collection_id"]} if context else {})
    }
    
    try:
        text = ""
        # Intentos específicos por tipo de archivo
        if "pdf" in mimetype:
            # Intento con PDFMiner como alternativa
            try:
                import pdfminer
                from pdfminer.high_level import extract_text as pdfminer_extract
                text = pdfminer_extract(file_path)
                logger.info("Extracción alternativa exitosa con PDFMiner", extra=error_context)
            except ImportError:
                logger.warning("PDFMiner no disponible para extracción alternativa", 
                            extra=error_context)
        elif "html" in mimetype or "xml" in mimetype:
            # Intentar con lxml como alternativa
            try:
                from lxml import etree
                parser = etree.HTMLParser()
                tree = etree.parse(file_path, parser)
                text = ' '.join(tree.xpath('//text()'))
                logger.info("Extracción alternativa exitosa con lxml", extra=error_context)
            except ImportError:
                logger.warning("lxml no disponible para extracción alternativa", 
                            extra=error_context)
                
        # Si ninguna alternativa funcionó, aplicar método genérico
        if not text or not text.strip():
            with open(file_path, 'rb') as f:
                content = f.read()
                # Intentar decodificar como texto si es posible
                encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'windows-1252']
                for encoding in encodings:
                    try:
                        text = content.decode(encoding)
                        if text.strip():
                            logger.info(f"Extracción genérica exitosa con encoding {encoding}", 
                                      extra=error_context)
                            break
                    except UnicodeDecodeError:
                        continue
        
        return text.strip() if text else ""
    except Exception as e:
        logger.error(f"Error en extracción alternativa: {str(e)}", 
                   extra=error_context, exc_info=True)
        # Usar formato de error estandarizado
        raise DocumentProcessingError(
            message=f"Failed to extract text using alternative methods: {str(e)}",
            details={"file_path": file_path, "error_code": ErrorCode.DOCUMENT_EXTRACTION_ERROR},
            context=error_context
        )

@with_context(tenant=True, validate_tenant=True)
@handle_errors(error_type="service", log_traceback=True)
async def process_text(
    text: str, 
    tenant_id: str,
    collection_id: str,
    metadata: Optional[Dict[str, Any]] = None,
    ctx: Context = None
) -> str:
    """
    Procesa texto plano directamente.

    Args:
        text: Texto a procesar
        tenant_id: ID del tenant
        collection_id: ID de la colección
        metadata: Metadatos adicionales

    Returns:
        str: Texto procesado
        
    Raises:
        DocumentProcessingError: Si hay un error en el procesamiento
    """
    if not text or not text.strip():
        raise ValidationError(
            message="Empty text provided for processing",
            details={"error_code": ErrorCode.EMPTY_TEXT}
        )
    
    return text

@with_context(tenant=True, validate_tenant=True)
@handle_errors(error_type="service", log_traceback=True)
async def process_file(
    file_path: str, 
    tenant_id: str,
    collection_id: str,
    metadata: Optional[Dict[str, Any]] = None,
    ctx: Context = None
) -> str:
    """
    Procesa un archivo local y extrae su contenido como texto.

    Args:
        file_path: Ruta al archivo local
        tenant_id: ID del tenant
        collection_id: ID de la colección
        metadata: Metadatos adicionales

    Returns:
        str: Texto extraído del archivo
        
    Raises:
        DocumentProcessingError: Si hay un error en la extracción
        ValidationError: Si el archivo no existe o no es válido
    """
    return await extract_text_from_file(file_path)

@with_context(tenant=True, validate_tenant=True)
@handle_errors(error_type="service", log_traceback=True)
async def process_file_from_storage(
    tenant_id: str,
    collection_id: str,
    file_key: str,
    metadata: Optional[Dict[str, Any]] = None,
    ctx: Context = None
) -> str:
    """
    Descarga un archivo desde el almacenamiento, lo procesa y extrae su contenido como texto.

    Args:
        tenant_id: ID del tenant
        collection_id: ID de la colección
        file_key: Clave del archivo en el almacenamiento
        metadata: Metadatos adicionales
        ctx: Contexto de la operación

    Returns:
        str: Texto extraído del archivo
        
    Raises:
        DocumentProcessingError: Si hay un error en la extracción
        ServiceError: Si hay un error en el acceso al almacenamiento
    """
    # Obtener contexto para errores
    error_context = {
        "service": "ingestion",
        "operation": "process_file_from_storage",
        "tenant_id": tenant_id,
        "collection_id": collection_id,
        "file_key": file_key
    }
    
    logger.info(f"Obteniendo archivo {file_key} del almacenamiento", extra=error_context)
    
    # Definir función para obtener de la base de datos
    async def fetch_from_storage(resource_id, tenant_id):
        logger.info(f"Obteniendo archivo {file_key} del almacenamiento", extra=error_context)
        try:
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                local_path = temp_file.name
                
            # Obtener archivo del storage
            file_data = await get_file_from_storage(file_key, local_path)
            if not file_data:
                return None
                
            file_data["local_path"] = local_path
            return file_data
        except Exception as e:
            logger.error(f"Error obteniendo archivo del almacenamiento: {str(e)}", 
                        extra=error_context, exc_info=True)
            return None
    
    # Definir función para verificar validez del archivo en caché
    async def validate_cached_data(cached_data):
        if not cached_data:
            return False
            
        local_path = cached_data.get("local_path")
        if not local_path:
            return False
            
        # Verificar que el archivo aún existe
        import os
        return os.path.exists(local_path)
    
    # Utilizar patrón de cache estandarizado
    try:
        resource_id = f"file:{file_key}"
        
        # Usar get_with_cache_aside siguiendo el patrón recomendado
        file_data, _ = await get_with_cache_aside(
            data_type="storage",
            resource_id=resource_id,
            tenant_id=tenant_id,
            collection_id=collection_id,
            fetch_from_db_func=fetch_from_storage,
            validate_cache_func=validate_cached_data,
            serialize_data=False  # Los datos ya están en formato serializable
        )
        
        if not file_data:
            raise ServiceError(
                message=f"Failed to retrieve file {file_key} from storage",
                details={"file_key": file_key},
                context=error_context
            )
            
        # Extraer texto
        local_path = file_data.get("local_path")
        mimetype = file_data.get("mimetype")
        logger.info(f"Procesando archivo {file_key} de tipo {mimetype}", extra=error_context)
        return await extract_text_from_file(local_path, mimetype)
    except Exception as e:
        logger.error(f"Error procesando archivo del almacenamiento: {str(e)}", 
                    extra=error_context, exc_info=True)
        
        # Limpiar archivo temporal si existe
        import os
        if 'local_path' in locals() and os.path.exists(local_path):
            try:
                os.unlink(local_path)
            except:
                pass
                
        raise

@with_context(tenant=True, validate_tenant=False)
@handle_errors(error_type="simple", log_traceback=False)
async def validate_file(file: UploadFile, ctx: Context = None) -> Dict[str, Any]:
    """
    Valida un archivo subido y verifica que cumple con los requisitos para ser procesado.
    
    Args:
        file: Archivo subido mediante FastAPI
        ctx: Contexto de la operación
        
    Returns:
        Dict[str, Any]: Información del archivo validado incluyendo mimetype, tamaño, etc.
        
    Raises:
        ValidationError: Si el archivo no es válido por alguna razón
    """
    # Obtener contexto para errores
    error_context = {
        "service": "ingestion",
        "operation": "validate_file",
        "file_name": file.filename
    }
    
    # Añadir información de contexto si está disponible
    if ctx:
        if ctx.has_tenant_id():
            error_context["tenant_id"] = ctx.get_tenant_id()
        if ctx.has_collection_id():
            error_context["collection_id"] = ctx.get_collection_id()
    
    # Verificar que se proporcionó un archivo
    if not file:
        raise ValidationError(
            message="No file provided",
            details={"error_code": ErrorCode.MISSING_FILE},
            context=error_context
        )
    
    # Verificar que el archivo tiene un nombre
    if not file.filename:
        raise ValidationError(
            message="File has no name",
            details={"error_code": ErrorCode.INVALID_FILENAME},
            context=error_context
        )
    
    # Detectar el tipo MIME
    mimetype = file.content_type
    
    # Verificar que el mimetype es soportado
    try:
        extractor_class = get_extractor_for_mimetype(mimetype)
    except DocumentProcessingError:
        raise ValidationError(
            message=f"Unsupported file type: {mimetype}",
            details={"mimetype": mimetype, "error_code": ErrorCode.UNSUPPORTED_FILE_TYPE},
            context=error_context
        )
    
    # Verificar el tamaño del archivo (límite configurable)
    settings = get_settings()
    max_file_size = settings.max_file_size_mb * 1024 * 1024  # Convertir MB a bytes
    
    # Intentar obtener el tamaño del archivo
    file_size = 0
    try:
        # Guardar la posición actual
        current_position = await file.tell()
        # Ir al final del archivo
        await file.seek(0, 2)  # 2 = SEEK_END
        # Obtener la posición final (tamaño)
        file_size = await file.tell()
        # Volver a la posición original
        await file.seek(current_position)
    except Exception as e:
        logger.warning(f"No se pudo determinar el tamaño del archivo: {str(e)}", 
                     extra=error_context)
    
    # Verificar si excede el tamaño máximo
    if file_size > max_file_size:
        max_size_mb = max_file_size / (1024 * 1024)
        actual_size_mb = file_size / (1024 * 1024)
        raise ValidationError(
            message=f"File too large. Maximum size is {max_size_mb:.1f} MB, but got {actual_size_mb:.1f} MB",
            details={
                "max_size_mb": max_size_mb,
                "file_size_mb": actual_size_mb,
                "error_code": ErrorCode.FILE_TOO_LARGE
            },
            context=error_context
        )
    
    # Devolver información del archivo
    return {
        "filename": file.filename,
        "mimetype": mimetype,
        "size": file_size,
        "extractor": extractor_class.__name__
    }