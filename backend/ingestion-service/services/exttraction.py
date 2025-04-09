"""
Funciones para extracción de texto de diferentes formatos de documento.
"""

import os
import logging
import mimetypes
from typing import Dict, Any, Optional

from common.errors import DocumentProcessingError
from config import get_extraction_config_for_mimetype

logger = logging.getLogger(__name__)

def detect_mimetype(file_path: str) -> str:
    """
    Detecta el tipo MIME de un archivo.
    
    Args:
        file_path: Ruta al archivo
        
    Returns:
        str: Tipo MIME del archivo
    """
    mime_type, _ = mimetypes.guess_type(file_path)
    
    if not mime_type:
        # Intentar detectar por extensión
        _, extension = os.path.splitext(file_path.lower())
        extension_map = {
            ".pdf": "application/pdf",
            ".txt": "text/plain",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc": "application/msword",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".xls": "application/vnd.ms-excel",
            ".csv": "text/csv",
            ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            ".ppt": "application/vnd.ms-powerpoint",
            ".md": "text/markdown",
            ".json": "application/json",
            ".html": "text/html",
            ".htm": "text/html"
        }
        mime_type = extension_map.get(extension, "application/octet-stream")
    
    return mime_type

async def extract_text_from_file(file_path: str, mimetype: Optional[str] = None) -> str:
    """
    Extrae texto de un archivo utilizando el extractor adecuado según su tipo MIME.
    
    Args:
        file_path: Ruta al archivo
        mimetype: Tipo MIME del archivo (opcional, se detecta si no se proporciona)
        
    Returns:
        str: Texto extraído del archivo
    """
    if not mimetype:
        mimetype = detect_mimetype(file_path)
    
    # Obtener configuración específica para este tipo MIME
    extraction_config = get_extraction_config_for_mimetype(mimetype)
    
    try:
        # Usar LlamaIndex SimpleDirectoryReader para extraer texto
        from llama_index.readers.file import SimpleDirectoryReader
        
        # El directorio debe existir para que SimpleDirectoryReader funcione
        file_dir = os.path.dirname(file_path)
        file_name = os.path.basename(file_path)
        
        # Configurar opciones según el tipo de documento
        file_extractor = SimpleDirectoryReader(
            input_dir=file_dir,
            file_extractor={
                # Configuraciones específicas para cada tipo
                ".pdf": extraction_config.get("pdf_parser", "default"),
                ".docx": "default",
                ".xlsx": "default",
                ".csv": "default"
            }
        )
        
        # Cargar solo el archivo específico
        documents = file_extractor.load_data(file=file_name)
        
        if not documents:
            logger.warning(f"No se pudo extraer texto del archivo {file_path}")
            return ""
        
        # Combinar texto de todos los documentos (puede haber múltiples páginas)
        combined_text = "\n\n".join([doc.text for doc in documents])
        
        return combined_text
    except Exception as e:
        logger.error(f"Error extrayendo texto de {file_path}: {str(e)}")
        # Intentar extraer con métodos específicos según el tipo MIME
        return await extract_with_specific_method(file_path, mimetype)

async def extract_with_specific_method(file_path: str, mimetype: str) -> str:
    """
    Intenta extraer texto con métodos específicos según el tipo MIME si LlamaIndex falla.
    
    Args:
        file_path: Ruta al archivo
        mimetype: Tipo MIME del archivo
        
    Returns:
        str: Texto extraído del archivo
    """
    try:
        # Implementaciones específicas para diferentes tipos MIME
        if mimetype == "application/pdf":
            return await extract_text_from_pdf(file_path)
        elif mimetype in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            return await extract_text_from_docx(file_path)
        elif mimetype in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel", "text/csv"]:
            return await extract_text_from_excel(file_path)
        elif mimetype in ["text/plain", "text/markdown", "application/json"]:
            # Archivos de texto plano
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        elif mimetype in ["text/html", "application/xhtml+xml"]:
            return await extract_text_from_html(file_path)
        else:
            logger.warning(f"No hay extractor específico para el tipo MIME: {mimetype}")
            # Intentar leer como texto plano
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            except:
                return ""
    except Exception as e:
        logger.error(f"Error en extracción específica para {mimetype}: {str(e)}")
        return ""

async def extract_text_from_pdf(file_path: str) -> str:
    """
    Extrae texto de un archivo PDF.
    
    Args:
        file_path: Ruta al archivo PDF
        
    Returns:
        str: Texto extraído del PDF
    """
    try:
        # Intentar primero con PyMuPDF (más rápido)
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text = ""
        
        for page in doc:
            text += page.get_text()
            text += "\n\n"
        
        return text
    except Exception as e:
        logger.warning(f"Error extrayendo con PyMuPDF: {str(e)}")
        
        try:
            # Intentar con pdfminer.six como alternativa
            from pdfminer.high_level import extract_text
            return extract_text(file_path)
        except Exception as e2:
            logger.error(f"Error extrayendo con pdfminer: {str(e2)}")
            
            try:
                # Último intento con PyPDF2
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                text = ""
                
                for page in reader.pages:
                    text += page.extract_text() + "\n\n"
                
                return text
            except Exception as e3:
                logger.error(f"Error extrayendo con PyPDF2: {str(e3)}")
                return ""

async def extract_text_from_docx(file_path: str) -> str:
    """
    Extrae texto de un archivo Word (.docx).
    
    Args:
        file_path: Ruta al archivo Word
        
    Returns:
        str: Texto extraído del documento
    """
    try:
        from docx import Document
        doc = Document(file_path)
        
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # También extraer texto de tablas
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error extrayendo texto de Word: {str(e)}")
        return ""

async def extract_text_from_excel(file_path: str) -> str:
    """
    Extrae texto de un archivo Excel (.xlsx, .xls) o CSV.
    
    Args:
        file_path: Ruta al archivo Excel o CSV
        
    Returns:
        str: Texto extraído del documento
    """
    try:
        import pandas as pd
        
        if file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
        else:  # Excel
            df = pd.read_excel(file_path)
        
        # Convertir cada hoja a texto
        text_parts = []
        
        # Añadir nombres de columnas
        text_parts.append(" | ".join(str(col) for col in df.columns))
        
        # Añadir filas
        for _, row in df.iterrows():
            text_parts.append(" | ".join(str(cell) for cell in row))
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Error extrayendo texto de Excel/CSV: {str(e)}")
        return ""

async def extract_text_from_html(file_path: str) -> str:
    """
    Extrae texto de un archivo HTML.
    
    Args:
        file_path: Ruta al archivo HTML
        
    Returns:
        str: Texto extraído del documento
    """
    try:
        from bs4 import BeautifulSoup
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        
        # Extraer texto visible y estructurado
        text = soup.get_text(separator='\n', strip=True)
        
        # También preservar enlaces importantes
        for link in soup.find_all('a', href=True):
            if link.text.strip():
                text += f"\nLink: {link.text.strip()} - {link['href']}"
        
        return text
    except Exception as e:
        logger.error(f"Error extrayendo texto de HTML: {str(e)}")
        return ""